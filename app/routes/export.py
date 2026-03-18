from flask import Blueprint, request, send_file, jsonify
import io
import re
from collections import defaultdict
import pandas as pd

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement # ضروري لاتجاه اليمين لليسار (RTL) في الجداول

from openpyxl.styles import Border, Side, Font, PatternFill, Alignment

from app.database import query_db

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

export_bp = Blueprint('export', __name__)

# ================== دوال مساعدة لتنسيق الوورد والإكسل ==================
def create_word_document_with_table(doc, title, headers, data):
    """دالة مساعدة لرسم جداول الوورد وتنسيقها مع دعم اليمين لليسار (RTL) الحقيقي"""
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # 1. جعل إطار الجدول نفسه من اليمين لليسار
    tblPr = table._tbl.tblPr
    if tblPr is not None:
        bidiVisual = OxmlElement('w:bidiVisual')
        tblPr.append(bidiVisual)
        
    # 2. دالة مساعدة داخلية لفرض اتجاه (RTL) على النصوص لعدم انقلاب الأقواس
    def set_rtl_run(cell):
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT # المحاذاة
            # فرض اتجاه الفقرة
            pPr = p._element.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')
            pPr.append(bidi)
            # فرض اتجاه الكلمات (لمنع انقلاب الأقواس)
            for run in p.runs:
                run.font.rtl = True 

    # كتابة رؤوس الجدول
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = str(header)
        set_rtl_run(hdr_cells[i])
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    # كتابة بيانات الجدول
    for row_data in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            set_rtl_run(row_cells[i])
            
    doc.add_page_break()

def process_and_format_sheet(writer, df, sheet_name, title=None, sheet_type=None):
    """دالة مساعدة لتنسيق جداول الإكسل باحترافية"""
    df.to_excel(writer, sheet_name=sheet_name)
    worksheet = writer.sheets[sheet_name]
    worksheet.sheet_view.rightToLeft = True # إجبار الإكسل على اليمين لليسار
    
    # 1. تعريف التنسيقات (الخطوط، الألوان، الحدود، والمحاذاة)
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
    center_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True) # التفاف النص ضروري جداً

    # 2. تنسيق الصف الأول (رؤوس الجدول: الأيام)
    for cell in worksheet[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center_alignment

    # 3. تنسيق باقي الخلايا (أوقات الحصص ومحتوى القاعات)
    for row in worksheet.iter_rows(min_row=2):
        for cell in row:
            cell.border = thin_border
            cell.alignment = center_alignment
            
    # 4. ضبط عرض الأعمدة لتناسب المحتوى
    worksheet.column_dimensions['A'].width = 18 # عرض عمود الوقت
    for col_idx in range(1, len(df.columns) + 1):
        col_letter = chr(65 + col_idx) # تحويل الرقم لحرف (B, C, D...)
        worksheet.column_dimensions[col_letter].width = 30 # عرض مناسب للقاعات المتعددة

# =====================================================================
# 1. تصدير جداول المستويات (Word)
# =====================================================================
@export_bp.route('/api/export/word/all-levels', methods=['POST'])
def export_all_levels_word():
    data = request.get_json()
    schedules_by_level, days, slots = data.get('schedule'), data.get('days', []), data.get('slots', [])
    if not all([schedules_by_level, days, slots]):
        return jsonify({"error": "بيانات التصدير غير كاملة"}), 400

    doc = Document()
    
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    margin_size = Cm(0.5)
    section.top_margin = margin_size
    section.bottom_margin = margin_size
    section.left_margin = margin_size
    section.right_margin = margin_size

    level_name_map = {"Bachelor 1": "ليسانس 1", "Bachelor 2": "ليسانس 2", "Bachelor 3": "ليسانس 3", "Master 1": "ماستر 1", "Master 2": "ماستر 2"}
    headers = ['الوقت'] + days

    for level, grid_data in schedules_by_level.items():
        processed_data = []
        for i, slot_name in enumerate(slots):
            row_content = [slot_name]
            for j in range(len(days)):
                cell_text = "\n".join([f"{lec.get('name', '')}\n{lec.get('teacher_name', '')}\n{lec.get('room', '')}".strip() for lec in grid_data[j][i]])
                row_content.append(cell_text)
            processed_data.append(row_content)
        
        sheet_name = level_name_map.get(level, level)
        create_word_document_with_table(doc, sheet_name, headers, processed_data)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name='جداول المستويات.docx')

# =====================================================================
# 2. تصدير جداول الأساتذة (Word)
# =====================================================================
@export_bp.route('/api/export/word/all-professors', methods=['POST'])
def export_all_professors_word():
    data = request.get_json()
    
    # ✨ --- التعديل تم في هذا السطر --- ✨
    schedules_by_prof = data.get('prof_schedules') or data.get('schedule')
    days = data.get('days', [])
    slots = data.get('slots', [])
    # ✨ ----------------------------- ✨

    if not all([schedules_by_prof, days, slots]):
        return jsonify({"error": "بيانات التصدير غير كاملة"}), 400

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    margin_size = Cm(0.5)
    section.top_margin = margin_size
    section.bottom_margin = margin_size
    section.left_margin = margin_size
    section.right_margin = margin_size

    level_name_map = {"Bachelor 1": "ليسانس 1", "Bachelor 2": "ليسانس 2", "Bachelor 3": "ليسانس 3", "Master 1": "ماستر 1", "Master 2": "ماستر 2"}
    headers = ['الوقت'] + days

    for prof_name, grid_data in sorted(schedules_by_prof.items()):
        processed_data = []
        for i, slot_name in enumerate(slots):
            row_content = [slot_name]
            for j in range(len(days)):
                cell_texts = [f"{lec.get('name', '')}\nالمستوى: {level_name_map.get(lec.get('level', ''), lec.get('level', ''))}\n{lec.get('room', '')}".strip() for lec in grid_data[j][i]]
                row_content.append("\n".join(cell_texts))
            processed_data.append(row_content)
        
        create_word_document_with_table(doc, prof_name, headers, processed_data)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name='جداول الأساتذة.docx')

# =====================================================================
# 3. تصدير القاعات الفارغة (Excel)
# =====================================================================
@export_bp.route('/api/export/free-rooms', methods=['POST'])
def export_free_rooms():
    data = request.get_json()
    free_rooms_grid = data.get('free_rooms') or data.get('schedule')
    days = data.get('days', [])
    slots = data.get('slots', [])
    
    if not all([free_rooms_grid, days, slots]): 
        return jsonify({"error": "بيانات التصدير غير كاملة"}), 400
    
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        processed_data = [["\n".join(free_rooms_grid[j][i]) for j in range(len(days))] for i in range(len(slots))]
        df = pd.DataFrame(processed_data, index=slots, columns=days)
        # استدعاء دالة التنسيق الاحترافية
        process_and_format_sheet(writer, df, 'القاعات الشاغرة')
    
    output.seek(0)
    return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', as_attachment=True, download_name='جدول_القاعات_الشاغرة.xlsx')
# =====================================================================
# 4. تصدير العبء البيداغوجي (Excel)
# =====================================================================
@export_bp.route('/api/export/teaching-load', methods=['GET'])
def export_teaching_load():
    try:
        # جلب المواد المسندة من قاعدة البيانات
        courses_raw = query_db('''
            SELECT c.id, c.name, t.name as teacher_name, group_concat(l.name, ',') as levels
            FROM courses c
            JOIN teachers t ON c.teacher_id = t.id
            LEFT JOIN course_levels cl ON c.id = cl.course_id
            LEFT JOIN levels l ON cl.level_id = l.id
            GROUP BY c.id
        ''')
        
        courses_by_teacher = defaultdict(list)
        for c in courses_raw:
            course_dict = dict(c)
            course_dict['levels'] = c['levels'].split(',') if c['levels'] else []
            courses_by_teacher[course_dict['teacher_name']].append(course_dict)

        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            workbook = writer.book
            if 'Sheet' in workbook.sheetnames:
                workbook.remove(workbook['Sheet'])
            worksheet = workbook.create_sheet('العبء البيداغوجي', 0)
            worksheet.sheet_view.rightToLeft = True

            thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
            header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell_alignment = Alignment(horizontal='right', vertical='top', wrap_text=True)
            merged_alignment = Alignment(horizontal='right', vertical='center', wrap_text=True)
            banded_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")

            headers = ['الرقم', 'اللقب', 'الاسم', 'الرتبة', 'الشهادة', 'تخصص الشهادة', 'المستوى', 'نوع المادة', 'الشعبة', 'التخصص', 'اسم المادة', 'القسم', 'الكلية', 'الحجم الساعي']
            for col_num, header_title in enumerate(headers, 1):
                cell = worksheet.cell(row=1, column=col_num)
                cell.value = header_title
                cell.font = header_font
                cell.fill = header_fill
                cell.border = thin_border
                cell.alignment = header_alignment

            worksheet.column_dimensions['A'].width = 5
            worksheet.column_dimensions['B'].width = 30
            for col_letter in ['C', 'D', 'E', 'F']: worksheet.column_dimensions[col_letter].width = 15
            for col_letter in ['G', 'H']: worksheet.column_dimensions[col_letter].width = 12
            for col_letter in ['I', 'J']: worksheet.column_dimensions[col_letter].width = 25
            worksheet.column_dimensions['K'].width = 45
            for col_letter in ['L', 'M', 'N']: worksheet.column_dimensions[col_letter].width = 15

            current_row = 2
            professor_number = 1
            for teacher_name in sorted(courses_by_teacher.keys()):
                courses = courses_by_teacher[teacher_name]
                total_rows_for_teacher = sum(len(c.get('levels', [])) for c in courses)
                if total_rows_for_teacher == 0: continue
                
                teacher_start_row = current_row
                is_banded_row = (professor_number % 2 == 0)

                for course in courses:
                    for level_name in course.get('levels', []):
                        course_name_original = course.get('name', '')
                        trimmed_course_name = re.sub(r'^\(.*\)\s*', '', course_name_original).strip()
                        level_abbr, course_type, division, specialization = '', 'تطبيق', '', ''
                        if 'سنة1 ليسانس' in level_name: level_abbr = 'ل1'
                        elif 'سنة2 ليسانس' in level_name: level_abbr = 'ل2'
                        elif 'سنة3 ليسانس' in level_name: level_abbr = 'ل3'
                        elif 'ماستر1' in level_name: level_abbr = 'م1'
                        elif 'ماستر2' in level_name: level_abbr = 'م2'
                        
                        if '[مح]' in course_name_original: course_type = 'محاضرة'
                        course_name_lower = course_name_original.lower()
                        
                        if level_abbr == 'ل1': division, specialization = 'جذع مشترك', 'جذع مشترك'
                        elif level_abbr == 'ل2':
                            if 'كل التخصصات' in course_name_lower: division, specialization = 'كل الشُّعب', 'كل التخصصات'
                            elif 'د.أدبية' in course_name_lower or 'د أدبية' in course_name_lower: division, specialization = 'دراسات أدبية', 'دراسات أدبية'
                            elif 'د.نقدية' in course_name_lower or 'د نقدية' in course_name_lower: division, specialization = 'دراسات نقدية', 'دراسات نقدية'
                            elif 'د.لغوية' in course_name_lower or 'د لغوية' in course_name_lower: division, specialization = 'دراسات لغوية', 'دراسات لغوية'
                        else:
                            if 'أدب عربي' in level_name or 'أدب عربي حديث ومعاصر' in level_name: division = 'دراسات أدبية'
                            elif 'لسانيات عامة' in level_name: division = 'دراسات لغوية'
                            elif 'نقد ومناهج' in level_name or 'نقد عربي قديم' in level_name: division = 'دراسات نقدية'
                            match = re.search(r'\((.*?)\)', level_name)
                            if match: specialization = match.group(1)

                        data_row = [level_abbr, course_type, division, specialization, trimmed_course_name]
                        for col_offset, value in enumerate(data_row):
                            worksheet.cell(row=current_row, column=7 + col_offset, value=value)
                        
                        current_row += 1

                end_row = current_row - 1
                
                for r in range(teacher_start_row, end_row + 1):
                    worksheet.row_dimensions[r].height = None
                    for c in range(1, 15):
                        cell = worksheet.cell(row=r, column=c)
                        if c >= 7 and c <= 11: cell.alignment = cell_alignment
                        else: cell.alignment = merged_alignment
                        cell.border = thin_border
                        if is_banded_row: cell.fill = banded_fill

                if total_rows_for_teacher > 0:
                    cell_a = worksheet.cell(row=teacher_start_row, column=1); cell_a.value = professor_number; cell_a.font = Font(bold=True)
                    cell_b = worksheet.cell(row=teacher_start_row, column=2); cell_b.value = teacher_name; cell_b.font = Font(bold=True)
                    if total_rows_for_teacher > 1:
                        worksheet.merge_cells(start_row=teacher_start_row, start_column=1, end_row=end_row, end_column=1)
                        worksheet.merge_cells(start_row=teacher_start_row, start_column=2, end_row=end_row, end_column=2)
                        for col in [3, 4, 5, 6, 12, 13, 14]:
                            worksheet.merge_cells(start_row=teacher_start_row, start_column=col, end_row=end_row, end_column=col)
                
                professor_number += 1

        output.seek(0)
        return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', as_attachment=True, download_name='العبء_البيداغوجي_للأساتذة.xlsx')

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"فشل إنشاء الملف: {e}"}), 500